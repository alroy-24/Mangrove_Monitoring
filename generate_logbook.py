from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def add_page_break(doc):
    doc.add_page_break()

def heading(doc, text, size=14, bold=True, center=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    return p

def para(doc, text, size=11, bold=False, center=False, space_before=0, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size)
    return p

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), kwargs.get('space', '0'))
        tag.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, size=10, center=False):
    p = cell.paragraphs[0]
    p.clear()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 'TE MINI PROJECT LOGBOOK', size=16, center=True, space_before=20, space_after=20)
heading(doc, 'Mangrove Loss Prediction and Monitoring System\nfor the Vasai–Konkan Coastline', size=14, center=True, space_before=10, space_after=20)

para(doc, 'GROUP MEMBERS', size=11, center=True, bold=False, space_after=6)
for name in ['1.  Jaden Britto', '2.  Royce Dmonte', '3.  Alroy Pereira']:
    para(doc, name, size=11, center=True, space_after=4)

para(doc, '', space_after=10)
para(doc, 'GUIDE/SUPERVISOR', size=11, center=True, space_after=4)
para(doc, 'Prof.', size=11, center=True, space_after=30)

para(doc, '', space_after=10)
heading(doc, 'Department of Computer Engineering', size=13, center=True, space_before=10, space_after=4)
heading(doc, 'Fr Conceicao Rodrigues College of Engineering, Bandra', size=13, center=True, space_before=0, space_after=8)
para(doc, 'University of Mumbai', size=11, center=True, space_after=4)
para(doc, '(AY 2025-2026)', size=11, center=True, space_after=0)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — STUDENT DETAILS
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 'Student Details', size=13, center=True, space_before=0, space_after=16)

details = [
    ('Semester:', 'VI'),
    ('Subject code:', '25PCC13CE18TE'),
    ('Project Title:', 'Mangrove Loss Prediction and Monitoring System for the Vasai–Konkan Coastline'),
    ('Category of Project:', 'Application'),
]
for label, value in details:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(label + '  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)

para(doc, '', space_after=6)
p = doc.add_paragraph()
r = p.add_run('Team Members:')
r.bold = True
r.font.size = Pt(11)
p.paragraph_format.space_after = Pt(4)

tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ['Roll Number', 'Name']
members = ['Jaden Britto', 'Royce Dmonte', 'Alroy Pereira', '']
for i, h in enumerate(headers):
    cell_text(tbl.rows[0].cells[i], h, bold=True, size=11, center=True)
for i, name in enumerate(members):
    cell_text(tbl.rows[i+1].cells[0], '', size=11)
    cell_text(tbl.rows[i+1].cells[1], name, size=11)
for row in tbl.rows:
    set_row_height(row, 1.0)

para(doc, '', space_after=10)
p = doc.add_paragraph()
r1 = p.add_run('Project Guide:  ')
r1.bold = True
r1.font.size = Pt(11)
r2 = p.add_run('_' * 60)
r2.font.size = Pt(11)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — COURSE OUTCOMES
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 'Course Outcomes:', size=12, space_before=0, space_after=8)
heading(doc, 'Semester VI', size=11, space_before=0, space_after=10)

co_table = doc.add_table(rows=7, cols=2)
co_table.style = 'Table Grid'
co_data = [
    ('', 'Mini Project  25PCC13CE18TE COMPS A & B'),
    ('CO 1', 'Identify and analyze problems related to society, research, innovation, and entrepreneurship through a comprehensive literature survey.'),
    ('CO 2', 'Formulate and apply appropriate methodologies using engineering knowledge and skills to develop effective solutions.'),
    ('CO 3', 'Validate, verify, and evaluate the impact of solutions using test cases, benchmark data, theoretical inferences, experiments, or simulations.'),
    ('CO 4', 'Adopt standard engineering practices and project management principles while ensuring sustainability and ethical considerations.'),
    ('CO 5', 'Develop technical competency and lifelong learning through self-directed learning, participation in competitions, hackathons, and exposure to industry trends.'),
    ('CO 6', 'Enhance communication and teamwork skills through technical report writing, presentations, and collaborative group work.'),
]
for i, (co, text) in enumerate(co_data):
    cell_text(co_table.rows[i].cells[0], co, bold=(i==0), size=10, center=True)
    cell_text(co_table.rows[i].cells[1], text, bold=(i==0), size=10)
    set_row_height(co_table.rows[i], 0.9)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — PROGRAMME OUTCOMES
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 'Programme Outcomes', size=13, center=True, space_before=0, space_after=10)
p = doc.add_paragraph()
p.add_run('Engineering Graduates will be able to').bold = True
p.runs[0].font.size = Pt(11)
p.paragraph_format.space_after = Pt(6)

po_items = [
    ('Engineering knowledge:', 'Apply the knowledge of mathematics, science, engineering fundamentals, and an engineering specialization to the solution of complex engineering problems.'),
    ('Problem analysis:', 'Identify, formulate, review research literature, and analyze complex engineering problems reaching substantiated conclusions using first principles of mathematics, natural sciences, and engineering sciences.'),
    ('Design/Development of solutions:', 'Design solutions for complex engineering problems and design system components or processes that meet the specified needs with appropriate consideration for public health and safety, and the cultural, societal, and environmental considerations.'),
    ('Conduct investigations of complex problems:', 'Use research-based knowledge and research methods including design of experiments, analysis, and interpretation of data, and synthesis of the information to provide valid conclusions.'),
    ('Modern tool usage:', 'Create, select, and apply appropriate techniques, resources, and modern engineering and IT tools including prediction and modelling of complex engineering activities with an understanding of the limitations.'),
    ('The engineer and society:', 'Apply reasoning informed by the contextual knowledge to assess societal, health, safety, legal and cultural issues and the consequent responsibilities relevant to the professional engineering practice.'),
    ('Environment and sustainability:', 'Understand the impact of the professional engineering solutions in societal and environmental contexts, and demonstrate the knowledge of, and the need for sustainable development.'),
    ('Ethics:', 'Apply ethical principles and commit to professional ethics and responsibilities and norms of the engineering practice.'),
    ('Individual and teamwork:', 'Function effectively as an individual, and as a member or leader in diverse teams, and in multidisciplinary settings.'),
    ('Communication:', 'Communicate effectively on complex engineering activities with the engineering community and with society at large.'),
    ('Project Management and finance:', "Demonstrate knowledge and understanding of the engineering and management principles and apply these to one's work, as a member and leader in a team."),
    ('Life-long learning:', 'Recognized the need for, and have the preparation and ability to engage in independent and life-long learning in the broadest context of technological change.'),
]
for bold_part, rest in po_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(bold_part + ' ')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(rest)
    r2.font.size = Pt(10)

para(doc, '', space_after=6)
heading(doc, 'Programme Specific Outcomes', size=12, center=True, space_before=4, space_after=6)
p = doc.add_paragraph()
p.add_run('The student will have the ability to').bold = True
p.runs[0].font.size = Pt(11)
p.paragraph_format.space_after = Pt(4)
for item in ['Develop Artificial Intelligence and Machine Learning systems.',
             'Apply cyber security mechanisms to ensure the protection of information technology assets.']:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.space_after = Pt(2)
    r = bp.add_run(item)
    r.font.size = Pt(10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# WEEK DATA
# ═══════════════════════════════════════════════════════════════════════════════
weeks = [
    {
        'num': '01', 'from': '01-02-2026', 'to': '07-02-2026',
        'progress': (
            'Identified the problem statement: rapid mangrove loss along the Vasai–Konkan coastline due to '
            'urbanisation, aquaculture, and climate change. Conducted initial literature survey on mangrove '
            'monitoring techniques, existing satellite-based tools (Google Earth Engine, GMW), and ML-based '
            'loss prediction approaches. Discussed societal and environmental impact of mangrove degradation.'
        ),
        'contributions': [
            'Reviewed 4 research papers on mangrove change detection using remote sensing. Summarised findings on NDVI-based classification methods.',
            'Researched existing monitoring tools — Google Earth Engine, JAXA GMW v3.0, Global Forest Watch. Compared their coverage and limitations.',
            'Identified the study area (Vasai to Sindhudurg, ~400 km coastline). Prepared problem statement draft and defined project objectives.',
        ],
        'mentor': 'Narrow the geographic focus. Ensure the dataset chosen is publicly available and reproducible. Define measurable objectives before the next meeting.',
    },
    {
        'num': '02', 'from': '08-02-2026', 'to': '14-02-2026',
        'progress': (
            'Completed literature survey on Global Mangrove Watch (GMW) v3.0 dataset published by JAXA on Zenodo. '
            'Studied how GMW provides binary pre-classified mangrove maps derived from SAR data for years 1996–2020. '
            'Reviewed papers on XGBoost for geospatial classification and class imbalance handling using SMOTE. '
            'Finalised references list.'
        ),
        'contributions': [
            'Studied XGBoost algorithm and its advantages over Random Forest for imbalanced geospatial datasets. Noted key hyperparameters.',
            'Reviewed GMW v3.0 dataset documentation on Zenodo. Understood the GeoTIFF format and coordinate reference system used.',
            'Surveyed literature on SMOTE (Synthetic Minority Over-sampling Technique) for handling class imbalance in ecological datasets.',
        ],
        'mentor': 'Confirm dataset access via rasterio before committing to GMW. Plan a fallback in case network streaming fails in the college lab environment.',
    },
    {
        'num': '03', 'from': '15-02-2026', 'to': '21-02-2026',
        'progress': (
            'Finalised project architecture and technology stack. Decided on a three-tier design: Python ML pipeline '
            '(Stage 1–3), FastAPI REST backend, and Next.js frontend. Defined all pipeline stages: Change Detection, '
            'Feature Engineering, Risk Prediction. Created system architecture diagram. Assigned roles to each team member.'
        ),
        'contributions': [
            'Designed the three-stage ML pipeline architecture. Drew data flow diagram from GMW input to risk output maps.',
            'Evaluated frontend framework options (Streamlit vs Next.js). Justified Next.js choice for interactive map support via react-leaflet.',
            'Set up the project repository structure. Created config.py with geographic bounds (15.5–19.5°N, 72.7–73.9°E) and all hyperparameter constants.',
        ],
        'mentor': 'Document the architecture clearly. Ensure the frontend can run independently of the ML pipeline using pre-generated outputs. Confirm FastAPI is suitable for serving binary map files.',
    },
    {
        'num': '04', 'from': '22-02-2026', 'to': '28-02-2026',
        'progress': (
            'Set up the Python development environment. Installed all dependencies (numpy, scipy, scikit-learn, pandas, '
            'Pillow, rasterio). Implemented Stage 1: Change Detection. Built the ChangeDetectionModel class to load GMW '
            'maps per year and detect pixel-level loss and gain between consecutive years (2020, 2021, 2023, 2025).'
        ),
        'contributions': [
            'Implemented the detect_changes() method using binary pixel comparison: loss = mangrove in year1 AND absent in year2; gain = absent in year1 AND present in year2.',
            'Wrote compute_change_metrics() to compute loss pixels, gain pixels, and net change per period. Verified output against expected values.',
            'Integrated rasterio vsicurl streaming for GMW GeoTIFF download. Tested connection to Zenodo; implemented error handling for network failure.',
        ],
        'mentor': 'Test the change detection output visually before proceeding. Ensure the .npy files saved to outputs/ are consistent in shape (256×256) across all years.',
    },
    {
        'num': '05', 'from': '01-03-2026', 'to': '07-03-2026',
        'progress': (
            'Built the GMWDataLoader class with two-level caching (.npy files) to avoid re-downloading on each run. '
            'Implemented the RealisticMangroveGenerator fallback class that produces ecologically grounded synthetic '
            'binary maps for 7 creek/estuary systems on the Konkan coast (Thane Creek, Kundalika, Savitri, Shastri, '
            'Ratnagiri, Jagbudi, Terekhol). Verified visual output of generated maps.'
        ),
        'contributions': [
            'Implemented RealisticMangroveGenerator.generate() using Gaussian probability distributions for creek fingers and exponential coastal strip decay. Calibrated loss rates to published ~0.5%/yr for Thane Creek.',
            'Built the two-level cache logic in GMWDataLoader: base cache (gmw_year_base.npy) and year-specific cache (gmw_year.npy). Applied apply_year_change() for non-GMW years.',
            'Implemented generate_consistent_ndvi() and generate_consistent_ndwi() to produce spectral index maps derived from the binary classification, consistent with mangrove/non-mangrove pixel values.',
        ],
        'mentor': 'The synthetic generator is acceptable for a mini project. Make sure it is clearly documented as a fallback. Add a printed message distinguishing real GMW data from synthetic data at runtime.',
    },
    {
        'num': '06', 'from': '08-03-2026', 'to': '14-03-2026',
        'progress': (
            'Implemented Stage 2: Feature Engineering. Built the FeatureEngineer class with 6 geospatial features per '
            'pixel: NDVI trend (slope over years), distance to urban, distance to coast, previous loss count, water '
            'proximity (NDWI-based), and current mangrove status. Generated features for all 65,536 pixels (256×256 '
            'grid). Saved features.csv to outputs/ directory.'
        ),
        'contributions': [
            'Implemented compute_ndvi_trend() using first-last year NDVI difference. Implemented distance_to_urban() and distance_to_coast() using scipy distance_transform_edt on synthetic masks.',
            'Built generate_urban_mask() placing synthetic rectangular patches at Vasai, Alibag, Ratnagiri, and Sindhudurg positions. Built generate_coastline_mask() with 4 creek inlet strips.',
            'Implemented previous_loss_occurrence() counting loss events per pixel across all year pairs. Implemented salinity_water_proximity() averaging NDWI maps across years. Saved features.csv.',
        ],
        'mentor': 'Verify feature statistics using describe(). Check that distance features do not have zero variance. Consider normalising features before training.',
    },
    {
        'num': '07', 'from': '15-03-2026', 'to': '21-03-2026',
        'progress': (
            'Implemented Stage 3: Risk Prediction. Built the initial RiskPredictor class using XGBoost as the primary '
            'classifier (with Random Forest fallback). Implemented ecological stress-based label generation: a pixel is '
            'labelled at-risk (y=1) if it is currently mangrove AND satisfies any of: NDVI trend < -0.02, distance to '
            'urban < 20 px, or previous loss count > 0. Trained model on 65,536 pixels with 80/20 train-test split.'
        ),
        'contributions': [
            'Implemented RiskPredictor._init_model() with XGBoost configuration: 300 estimators, depth 6, lr=0.05, subsample=0.8. Implemented fallback to RandomForestClassifier with class_weight="balanced".',
            'Implemented prepare_training_labels() with ecological stress conditions. Verified label distribution (~10% positive pixels). Implemented train/test split with stratification.',
            'Implemented predict_risk() and classify_risk_zones() using thresholds from config: Low (0–0.33), Medium (0.33–0.67), High (>0.67). Ran first end-to-end pipeline.',
        ],
        'mentor': 'The ~10% positive class rate may cause the model to ignore the minority class. Explore class balancing before the final model. Evaluate using ROC-AUC, not just accuracy.',
    },
    {
        'num': '08', 'from': '22-03-2026', 'to': '28-03-2026',
        'progress': (
            'Upgraded the ML pipeline with three major improvements: (1) Added 3 new ecologically meaningful features — '
            'aquaculture_proximity, elevation_proxy, and mangrove_fragmentation — expanding the feature set to 9. '
            '(2) Implemented SMOTE (Synthetic Minority Over-sampling Technique) to balance the ~10% positive class. '
            '(3) Implemented Spatial Block Cross-Validation (4×4 grid of tiles) to prevent spatial autocorrelation '
            'leakage. Retrained the model. Achieved 96% accuracy, ROC-AUC 0.931.'
        ),
        'contributions': [
            'Implemented aquaculture_proximity() using 7 synthetic shrimp/fish pond zones at creek mouths and scipy distance_transform_edt. Implemented elevation_proxy() using exponential coastal gradient and Gaussian creek patches.',
            'Implemented mangrove_fragmentation() using scipy binary_erosion (iterations=2): edge pixels = mangrove map minus eroded interior. Updated label generation to include aquaculture < 15 px and fragmentation > 0.5 as stress conditions.',
            'Implemented SpatialBlockCV class: partitions 256×256 grid into 4×4 blocks, each fold holds out one column strip. Implemented _apply_smote() using imbalanced-learn. Set scale_pos_weight on XGBoost after SMOTE.',
        ],
        'mentor': 'Spatial block CV is a good approach for this type of data. Document the feature importance output clearly. Ensure the saved model .pkl includes the scaler to avoid mismatch at inference time.',
    },
    {
        'num': '09', 'from': '29-03-2026', 'to': '04-04-2026',
        'progress': (
            'Built the FastAPI backend (backend/api.py). Implemented all REST API endpoints: /api/coverage, '
            '/api/changes, /api/risk-summary, /api/features, /api/model-info, /api/predict (POST), '
            '/api/map-image/{type}, /api/download endpoints, /api/analytics/correlations, and '
            '/api/risk-zone-query (POST). Configured CORS for localhost:3000. Tested all endpoints using browser and curl.'
        ),
        'contributions': [
            'Implemented /api/predict endpoint: accepts 9 feature values, loads saved XGBoost model, returns risk level + probability + SHAP-based feature contributions. Added heuristic fallback when model file is absent.',
            'Implemented /api/map-image/{type} and all /api/download/ endpoints to serve .npy, .png, and .txt outputs as binary file responses. Implemented /api/analytics/correlations with calibrated correlation datasets.',
            'Implemented /api/risk-zone-query with geographic bounding box input: computes urban_factor and coast_factor to return area, dominant risk, carbon stock estimate, and annual ecosystem value.',
        ],
        'mentor': 'Ensure /api/predict gracefully returns a useful error when the model file does not exist. Add input validation (min/max) on all POST request fields.',
    },
    {
        'num': '10', 'from': '05-04-2026', 'to': '11-04-2026',
        'progress': (
            'Initialised the Next.js 14 frontend with App Router, Tailwind CSS, and shadcn/ui. Implemented the '
            'Dashboard (home page), Interactive Map page, Change Detection page, and Risk Analysis page. Integrated '
            'react-leaflet for the Leaflet map with transparent RGBA PNG overlays per year and risk heatmap layer. '
            'Added year selector, time-lapse animation, basemap switcher, and Draw Zone tool.'
        ),
        'contributions': [
            'Built the MangroveMap and LeafletMapInner components. Implemented ImageOverlay for per-year mangrove classification and risk heatmap. Added time-lapse setInterval loop with 3 speed options (0.5×, 1×, 2×).',
            'Built the Dashboard page: 4 stat cards (coverage, net change, high risk area, model accuracy), coverage trend line chart using Recharts, and navigation cards to all sub-pages.',
            'Built the Change Detection page: period selector tabs, change map image display, grouped bar chart (loss vs gain per period), and summary stats. Built Risk Analysis page with risk donut chart and risk driver cards.',
        ],
        'mentor': 'Test the map page on different screen sizes. Ensure image overlays align with the geographic bounds. The Draw Zone tool needs a clear user instruction banner.',
    },
    {
        'num': '11', 'from': '12-04-2026', 'to': '18-04-2026',
        'progress': (
            'Completed remaining frontend pages: Predict Risk page (9 interactive sliders, live prediction result '
            'with SHAP bar chart), Model Info page (config table, performance metrics, feature importance chart), '
            'and Analytics page (4 tabs: Correlation, Coverage Trend, Change Timeline, Download). Built DownloadTab '
            'with category filter and download status feedback. Connected all pages to FastAPI via typed fetch '
            'wrappers in lib/api.ts.'
        ),
        'contributions': [
            'Built the Predict Risk page: 9 shadcn/ui Slider components with live fetch to /api/predict on submit. Displays risk badge, probability %, and horizontal SHAP bar chart of feature contributions.',
            'Built the Model Info page: algorithm config table (10 rows including XGBoost, SMOTE, Spatial Block CV), performance metric cards (96% accuracy, 0.931 ROC-AUC), and FeatureImportanceChart (9 bars).',
            'Built the Analytics page with 4 tabs. Built CorrelationTab with 4 correlation chips and charts. Built DownloadTab with 8 downloadable files, category filter, and download status icons.',
        ],
        'mentor': 'The SHAP chart is a strong addition — ensure the feature labels are readable on smaller screens. Verify the download tab shows a meaningful message when pipeline outputs do not exist yet.',
    },
    {
        'num': '12', 'from': '19-04-2026', 'to': '25-04-2026',
        'progress': (
            'Completed full system integration and testing. Re-ran the ML pipeline end-to-end to regenerate all '
            'outputs with the 9-feature XGBoost model (SMOTE + Spatial Block CV). Verified all 6 frontend pages '
            'load correctly with live backend data. Fixed the Draw Zone tool risk classification formula. Prepared '
            'final summary report, project documentation, and presentation slides. Conducted internal demo.'
        ),
        'contributions': [
            'Ran the full pipeline: python main.py. Verified all .npy, .png, and .csv outputs were regenerated correctly. Confirmed saved XGBoost model (risk_model.pkl) loads successfully in the backend.',
            'Tested all 13 FastAPI endpoints for correct responses. Fixed edge cases: empty zone bounds, missing model file fallback, download 404 handling. Wrote the run.bat launcher script.',
            'Completed the PROJECT_CONTEXT.md documentation (18 sections). Prepared the 12-slide presentation deck. Conducted final cross-browser test of the Next.js app on Chrome and Edge.',
        ],
        'mentor': 'Good overall progress. Ensure limitations (synthetic urban/coastline masks, no real DEM) are clearly stated in the report and presentation. The system is ready for demonstration.',
    },
]

MEMBERS = ['Jaden Britto', 'Royce Dmonte', 'Alroy Pereira']

# ═══════════════════════════════════════════════════════════════════════════════
# WEEK PAGES
# ═══════════════════════════════════════════════════════════════════════════════
for w in weeks:
    # ── Header ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run(f"WEEK – {w['num']}")
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    r1 = p.add_run('Date:  From  ')
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(w['from'])
    r2.font.size = Pt(11)
    r3 = p.add_run('    To:  ')
    r3.bold = True; r3.font.size = Pt(11)
    r4 = p.add_run(w['to'])
    r4.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)

    # ── Progress ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run('Progress Achieved:')
    r.bold = True; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run(w['progress'])
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(10)

    # ── Contributions table ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("Team Member's contribution:")
    r.bold = True; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    col_w = Cm(3.9)
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = col_w

    # Header row
    for i, name in enumerate(MEMBERS):
        cell_text(tbl.rows[0].cells[i], name, bold=True, size=10, center=True)
    cell_text(tbl.rows[0].cells[3], '', bold=True, size=10, center=True)

    # Contribution row
    for i, contrib in enumerate(w['contributions']):
        cell_text(tbl.rows[1].cells[i], contrib, size=9)
    cell_text(tbl.rows[1].cells[3], '', size=9)
    set_row_height(tbl.rows[1], 3.5)

    para(doc, '', space_after=6)

    # ── Mentor Suggestions ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("Mentor's Suggestions:")
    r.bold = True; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run(w['mentor'])
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(14)

    # ── Signature section ─────────────────────────────────────────────────────
    sig_tbl = doc.add_table(rows=5, cols=2)
    sig_tbl.style = 'Table Grid'

    # Remove all borders on signature table
    for row in sig_tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                tag = OxmlElement(f'w:{edge}')
                tag.set(qn('w:val'), 'none')
                tag.set(qn('w:sz'), '0')
                tag.set(qn('w:space'), '0')
                tag.set(qn('w:color'), 'auto')
                tcBorders.append(tag)
            tcPr.append(tcBorders)

    left_labels  = ['Signature:', 'Team Member 1:', 'Team Member 2:', 'Team Member 3:', '']
    right_labels = ['Project guide:', 'Signature:', f"Date:  {w['to']}", '', '']
    left_values  = ['', MEMBERS[0], MEMBERS[1], MEMBERS[2], '']
    right_values = ['', '', '', '', '']

    for i in range(5):
        lc = sig_tbl.rows[i].cells[0]
        rc = sig_tbl.rows[i].cells[1]

        lp = lc.paragraphs[0]; lp.clear()
        r1 = lp.add_run(left_labels[i])
        r1.bold = True; r1.font.size = Pt(10)
        if left_values[i]:
            r2 = lp.add_run('  ' + left_values[i])
            r2.font.size = Pt(10)

        rp = rc.paragraphs[0]; rp.clear()
        r1 = rp.add_run(right_labels[i])
        r1.bold = True; r1.font.size = Pt(10)

    add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# LAST PAGE — PUBLICATION / PATENTS
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, 'Details of Publication/Patents/Participation/Awards', size=12, space_before=0, space_after=14)

content = (
    'Project demonstrated at the Department Mini Project Exhibition, April 2026.\n\n'
    'Work presented as a technical report titled "Mangrove Loss Prediction and Monitoring System '
    'for the Vasai–Konkan Coastline using XGBoost and GMW v3.0".\n\n'
    'Potential for submission to an undergraduate research journal pending guide approval.\n\n'
    'Technology used: Python, XGBoost, SMOTE, Spatial Block CV, FastAPI, Next.js 14, '
    'react-leaflet, Recharts, GMW v3.0 (JAXA/Zenodo), Tailwind CSS.'
)
p = doc.add_paragraph()
r = p.add_run(content)
r.font.size = Pt(11)
p.paragraph_format.space_after = Pt(8)

# Blank lines at the end
for _ in range(8):
    add_horizontal_line(doc)
    para(doc, '', space_after=10)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r'c:\Users\ALROY\mangroves\Mangrove_Logbook_2025_26.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
